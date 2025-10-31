import os
from telegram import Update, ReplyKeyboardMarkup, InlineQueryResultArticle, InputTextMessageContent
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, InlineQueryHandler, ContextTypes, filters
from translator import is_kiril, kiril_to_lotin, lotin_to_kiril
from utils import add_user_stat, load_stats
from dotenv import load_dotenv
from uuid import uuid4
from docx import Document
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas

load_dotenv()
TOKEN = os.getenv("TOKEN")

menu = ReplyKeyboardMarkup([
    ["🔤 Matn tarjima", "📄 Fayl tarjima"],
    ["📊 Statistikalar", "ℹ️ Yordam"]
], resize_keyboard=True)

# 💬 /start
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 Assalomu alaykum!\nMen Kiril ↔ Lotin tarjima botiman.\nQuyidagilardan birini tanlang:",
        reply_markup=menu
    )

# ℹ️ /help (va Yordam tugmasi)
async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "ℹ️ <b>Yordam:</b>\n\n"
        "🔤 <b>Matn tarjima:</b> Oddiy matn yuboring — avtomatik tarjima qilinadi.\n\n"
        "📄 <b>Fayl tarjima:</b> .docx, .pdf yoki .txt fayl yuboring — tarjima qilib o‘sha formatda qaytaraman.\n\n"
        "📊 <b>Statistikalar:</b> Siz nechta matn va fayl tarjima qilganingizni ko‘rishingiz mumkin.\n\n"
        "📢 <b>Kanal:</b> @AsadbekAzamatovich",
        parse_mode="HTML"
    )

# 📊 Statistika
async def stats(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    stat = load_stats().get(str(user_id), {"texts": 0, "files": 0})
    await update.message.reply_text(
        f"📊 Sizning statistikangiz:\n\n"
        f"🔤 Matn tarjimalari: {stat['texts']}\n"
        f"📄 Fayl tarjimalari: {stat['files']}"
    )

# 🔘 Tugmalarni tushuntirish
async def menu_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    if text == "🔤 Matn tarjima":
        await update.message.reply_text("✍️ Siz oddiy matn yuborishingiz mumkin. Men uni avtomatik tarjima qilaman.")
    elif text == "📄 Fayl tarjima":
        await update.message.reply_text("📎 Siz Word (.docx), PDF yoki TXT fayl yuborishingiz mumkin. Men uni tarjima qilib o‘sha formatda qaytaraman.")
    elif text == "📊 Statistikalar":
        await stats(update, context)
    elif text == "ℹ️ Yordam":
        await help_command(update, context)

# 📝 Matn tarjimasi
async def translate_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    if text in ["🔤 Matn tarjima", "📄 Fayl tarjima", "📊 Statistikalar", "ℹ️ Yordam"]:
        return
    translated = kiril_to_lotin(text) if is_kiril(text) else lotin_to_kiril(text)
    add_user_stat(update.message.from_user.id, text_count=1)
    await update.message.reply_text(translated)

# 📄 Fayl tarjimasi
async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    file = await update.message.document.get_file()
    filename = update.message.document.file_name
    ext = filename.split('.')[-1].lower()

    local_path = f"temp_input.{ext}"
    await file.download_to_drive(local_path)

    text = ""
    if ext == "txt":
        text = open(local_path, 'r', encoding='utf-8', errors='ignore').read()
    elif ext == "docx":
        doc = Document(local_path)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif ext == "pdf":
        reader = PdfReader(local_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    else:
        await update.message.reply_text("❌ Faqat .txt, .docx yoki .pdf fayl yuboring.")
        return

    translated = kiril_to_lotin(text) if is_kiril(text) else lotin_to_kiril(text)
    add_user_stat(update.message.from_user.id, file_count=1)

    if ext == "txt":
        out_path = "translated.txt"
        open(out_path, "w", encoding="utf-8").write(translated)
    elif ext == "docx":
        out_path = "translated.docx"
        doc = Document()
        for line in translated.split("\n"):
            doc.add_paragraph(line)
        doc.save(out_path)
    else:
        out_path = "translated.pdf"
        c = canvas.Canvas(out_path)
        c.setFont("Helvetica", 11)
        y = 800
        for line in translated.split("\n"):
            if y < 50:
                c.showPage()
                y = 800
            c.drawString(40, y, line)
            y -= 15
        c.save()

    await update.message.reply_document(open(out_path, "rb"))
    os.remove(local_path)
    os.remove(out_path)

# 🔍 Inline tarjima
async def inline_translate(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.inline_query.query
    if not query:
        return
    translated = kiril_to_lotin(query) if is_kiril(query) else lotin_to_kiril(query)
    results = [
        InlineQueryResultArticle(
            id=str(uuid4()),
            title="Tarjima natijasi",
            input_message_content=InputTextMessageContent(translated)
        )
    ]
    await update.inline_query.answer(results)

# 🚀 Ishga tushurish
if __name__ == "__main__":
    app = ApplicationBuilder().token(TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(MessageHandler(filters.Regex("^(🔤 Matn tarjima|📄 Fayl tarjima|📊 Statistikalar|ℹ️ Yordam)$"), menu_handler))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, translate_text))
    app.add_handler(InlineQueryHandler(inline_translate))

    print("✅ Bot ishga tushdi...")
    app.run_polling()
